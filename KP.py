import os
from tkinter import *
from tkinter import filedialog, messagebox, colorchooser
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class ImageToWordApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Image to Word Converter")

        self.label = Label(root, text="Выберите папку с изображениями:")
        self.label.pack(pady=10)

        self.choose_button = Button(root, text="Выбрать папку", command=self.choose_directory)
        self.choose_button.pack(pady=5)

        self.convert_button = Button(root, text="Создать документ Word", command=self.convert_to_word)
        self.convert_button.pack(pady=10)

    def choose_directory(self):
        self.directory_path = filedialog.askdirectory()
        if self.directory_path:
            messagebox.showinfo("Папка выбрана", f"Выбрана папка: {self.directory_path}")

    def convert_to_word(self):
        if not hasattr(self, 'directory_path'):
            messagebox.showerror("Ошибка", "Пожалуйста, выберите папку с изображениями.")
            return

        output_docx = filedialog.asksaveasfilename(defaultextension=".docx",
                                                  filetypes=[("Word документы", "*.docx")],
                                                  initialfile="output_document.docx",
                                                  title="Сохранить документ Word как")

        if not output_docx:
            return

        color_info_1 = colorchooser.askcolor(title="Выберите цвет для границы аэрофотосъемки")
        if not color_info_1 or not color_info_1[1]:
            return

        color_info_2 = colorchooser.askcolor(title="Выберите цвет для границ районов работ")
        if not color_info_2 or not color_info_2[1]:
            return

        text_color_1 = colorchooser.askcolor(title="Выберите цвет текста для границы аэрофотосъемки")
        if not text_color_1 or not text_color_1[1]:
            return

        text_color_2 = colorchooser.askcolor(title="Выберите цвет текста для границ районов работ")
        if not text_color_2 or not text_color_2[1]:
            return

        legend_line_color_1 = color_info_1[1]
        legend_line_color_2 = color_info_2[1]
        legend_text_color_1 = text_color_1[1]
        legend_text_color_2 = text_color_2[1]

        try:
            doc = Document()

            # Установка шрифта и размера для всего документа
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            # Установка шрифта и размера для заголовков
            heading_styles = ['Heading 1', 'Heading 2', 'Heading 3']
            for style_name in heading_styles:
                heading_style = doc.styles[style_name]
                heading_font = heading_style.font
                heading_font.name = 'Times New Roman'
                heading_font.size = Pt(12)
                heading_font.bold = True
                heading_font.color.rgb = RGBColor(0, 0, 0)

            # Добавляем первую страницу по шаблону
            heading = doc.add_heading('ПАСПОРТ АЭРОФОТОСЪЕМКИ', level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc.add_paragraph("Исполнитель: Филиал ППК «Роскадастр» «Аэрогеодезия»\n"
                              "Календарный период выполнения аэрофотосъемки: 20.04.2024 – 16.06.2024\n"
                              "Шифр объекта: 10.15.4.3600.ГП54М2 (Тамбовская область)\n"
                              "Фактическая площадь аэрофотосъемки, кв. км:\n"
                              "Размер пикселя на местности, м (для цифровых АС), м: 0.06 – 0.18\n"
                              "Высота фотографирования, м: 500 – 1400 м\n"
                              "Вид аэрофотосъемки (площадная, линейная): площадная\n"
                              "Ориентация маршрутов (широтная, меридиональная, заданная): заданная проектом\n"
                              "Тип (аналоговая, цифровая), спецификация и номер АС: Sony DSC-RX1 RM2: SN 9679324, SN 9679379; SN 110000078, SN 110000072; Sony DSC-RX1: SN 7452967; SN 7452934, SN 7452937\n"
                              "Тип и серийный номер объектива (если объектив съемный, заменяемый): -\n"
                              "Фокусное расстояние, мм: 35\n"
                              "Размеры аэрофотоснимка (Nx-Ny), выраженные в пикселях (для цифровой АС), пикс: Nx = 7952; Ny = 5304; Nx = 6000; Ny = 4000\n"
                              "Физический размер пикселя, мм (для цифровой АС), мм: 0.0044 на 0.0044\n"
                              "Ориентация системы координат снимка (для кадровой АС) относительно направления полета (ось X вперед, назад, влево, вправо): вправо\n"
                              "Аэрофотоустановка (гироплатформа): -\n"
                              "Спектральная характеристика аэрофотоснимков (RGB, PAN, прочее): RGB\n"
                              "Дополнительная аппаратура: Система определения положения и ориентации, тип, номер: Topcon b111 SN VFBHR18440138, SN 19370025; SN 18160094, SN VFBHR19050021, SN VFBHR19050027, SN 19420257\n"
                              "Тип воздушного судна: БВС Геоскан 201: борт. №20366, №20396, №20640, №20641, №20365, №20395, №20388, №20389\n"
                              "Число маршрутов: 4\n"
                              "Общее число снимков: 203")

            images = os.listdir(self.directory_path)
            image_number = 1

            for image_name in images:
                if image_name.endswith('.jpg') or image_name.endswith('.png'):
                    doc.add_page_break()

                    # Добавляем заголовок перед снимком
                    heading = doc.add_heading(f'Схема покрытия территории аэрофотоснимками\n ', level=2)
                    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    heading = doc.add_heading(f'{image_number}. {image_name}\n', level=2)
                    heading.runs[0].font.bold = False
                    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    img_path = os.path.join(self.directory_path, image_name)
                    # Установить размер изображения
                    doc.add_picture(img_path, width=Inches(5.5))

                    # Центрирование изображения
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # Добавляем легенду под изображением
                    legend_paragraph = doc.add_paragraph()
                    legend_run = legend_paragraph.add_run("Условные обозначения:\n")
                    legend_run.font.color.rgb = RGBColor(0, 0, 0)
                    legend_run = legend_paragraph.add_run("------------- ")
                    legend_run.font.color.rgb = self.hex_to_rgb(legend_line_color_1)
                    legend_run = legend_paragraph.add_run("― граница аэрофотосъемки\n")
                    legend_run.font.color.rgb = self.hex_to_rgb(legend_text_color_1)
                    legend_run = legend_paragraph.add_run("------------- ")
                    legend_run.font.color.rgb = self.hex_to_rgb(legend_line_color_2)
                    legend_run = legend_paragraph.add_run("― границы районов работ")
                    legend_run.font.color.rgb = self.hex_to_rgb(legend_text_color_2)
                    legend_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                    image_number += 1

            # Добавляем завершающую страницу с описью
            doc.add_page_break()

            # Добавляем описание материалов аэрофотосъемки
            doc.add_paragraph("Опись материалов аэрофотосъемки на территории Тамбовской области, выполненной в 2024 году.")
            doc.add_paragraph(f"Шифр объекта: 10.15.4.3600.ГП54М2 (Тамбовская область)\n"
                              f"Материалы аэрофотосъемки в электронном виде расположены на HDD-диске, который содержит папку с наименованием региона, в котором была выполнена АФС. Аэрофотоснимки находятся в папке «Photo», в формате *JPG.\n"
                              f"Общий объем содержимого диска – 3.68 ГБ.\n"
                              f"Опись составил:\n"
                              f"Начальник аэрофотосъемочной группы:                               Рыбко М.А.")

            # Сохраняем документ
            doc.save(output_docx)
            messagebox.showinfo("Успех", f"Документ успешно сохранен как {output_docx}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Произошла ошибка при создании документа: {e}")

    def hex_to_rgb(self, hex_color):
        hex_color = hex_color.lstrip('#')
        return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

if __name__ == "__main__":
    root = Tk()
    app = ImageToWordApp(root)
    root.mainloop()
